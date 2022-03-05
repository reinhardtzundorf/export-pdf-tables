# export.py
# =========

# Setup 
# -----

import os
import sys
import docx
import pandas as pd

file_name = sys.argv[1]
file_type = sys.argv[2]
file_number = sys.argv[3]

# Get Data  
# --------

dfs = pd.read_csv(f"tables/{file_type}/{file_name}/{file_number}.{file_type}")

# Generate Document
# -----------------

doc = docx.Document() 
doc.save("test.docx")
table = doc.add_table(rows=dfs.shape[0], cols=dfs.shape[1])
table_cells = table._cells
for i in range(dfs.shape[0]):
    for j in range(dfs.shape[1]):
        table_cells[j + i * dfs.shape[1]].text =  str(dfs.values[i][j])

# Save Document
# -------------

doc = docx.Document() 
table = doc.add_table(rows=dfs.shape[0], cols=dfs.shape[1])
table_cells = table._cells
for i in range(dfs.shape[0]):
    for j in range(dfs.shape[1]):
        table_cells[j + i * dfs.shape[1]].text =  str(dfs.values[i][j])
if os.path.isdir(f"tables/docx/{file_name}") == False:
	os.makedirs(f"tables/docx/{file_name}")
doc.save(f"tables/docx/{file_name}/{file_number}.docx")