# wsgi.py

from flask import Flask, request
from pathlib import Path
import requests
import os
import sys
import docx
import time
import tabula
import pandas as pd


app = Flask(__name__,
            static_url_path='/assets', 
            static_folder='assets',
            template_folder='templates')


@app.route("/", methods=["GET"])
def form_example():
    return open('src/templates/index.html').read()


@app.route("/", methods=["POST"])
def form_submit():
    file_name = request.form.get("file_name")
    file_type = request.form.get("file_type")
    dry_run = request.form.get("dry_run")
    
    # Upload file.
    uploaded_file = upload_file()

    # Extract the tables to CSV.
    count = extract_pdf_tables_to_csv(uploaded_file, 'test4')

    # Convert the CSV files to docx files for each table.
    convert_outcome = convert_csv_tables_to_docx('test4', count, 'csv')

    # Return outcome.
    return convert_outcome

    # return f'upload_file_name{uploaded_file}'
    # return f'file_name: {file_name}<br>dry_run: {dry_run}<br>file_type: {file_type}<br>'


# Upload File
# ===========
def upload_file():
   if request.method == 'POST':
      f = request.files['file_pdf']
      f.save(f'uploads/{f.filename}')
      return f.filename


# Extract PDF Tables To CSV
# =========================
def extract_pdf_tables_to_csv(pdf_path, file_name): 
   
   # Report
   report = ''

   # Counter
   count = 0

   # Path to PDF file to extract tables from.
   pdf_path = f"uploads/{pdf_path}"

   # Path to directory for JSON table data.
   json_table_path = "tables/json/"

   # Path to directory for CSV table data.
   csv_table_path = "tables/csv/"

   # Extract Table(s).

   # The following function parses the PDF file using the tabula library's 
   # read_pdf function. The function iterates over the contents of the file 
   # and sets any tables it finds as a DataFrame object.

   dfs = tabula.read_pdf(pdf_path, pages='all')

   # Output length.
   print(len(dfs))

   # Output the first DataFrame.
   print(dfs[0])

   # Check if directory exists.
   if os.path.isdir(f"{csv_table_path}{file_name}") == False:
      os.makedirs(f"{csv_table_path}{file_name}")

   # Save each table to CSV file.
   for i in range(len(dfs)):
       dfs[i].to_csv(f"{csv_table_path}{file_name}/{i}.csv")
       count = count + 1

   return count


# Convert CSV Tables To Docx
# ==========================
def convert_csv_tables_to_docx(csv_prefix_name, num_of_tables, file_type): 

   # Loop Over Tables
   # ----------------
   for i in range(num_of_tables):

      file_name = csv_prefix_name
      file_type = file_type
      file_number = i

      # Get Data  
      # --------
      dfs = pd.read_csv(f"tables/{file_type}/{file_name}/{file_number}.{file_type}")

      # Generate Document
      # -----------------
      doc = docx.Document() 
      doc.save("test.docx")
      table = doc.add_table(rows=dfs.shape[0], cols=dfs.shape[1])
      table_cells = table._cells
      for i in range(dfs.shape[0]):
          for j in range(dfs.shape[1]):
              table_cells[j + i * dfs.shape[1]].text =  str(dfs.values[i][j])

      # Save Document
      # -------------
      doc = docx.Document() 
      table = doc.add_table(rows=dfs.shape[0], cols=dfs.shape[1])
      table_cells = table._cells
      for i in range(dfs.shape[0]):
          for j in range(dfs.shape[1]):
              table_cells[j + i * dfs.shape[1]].text =  str(dfs.values[i][j])
      if os.path.isdir(f"tables/docx/{file_name}") == False:
         os.makedirs(f"tables/docx/{file_name}")
      doc.save(f"tables/docx/{file_name}/{file_number}.docx")

   return 'success'
